# %%
import pandas as pd
from docx import Document
from docx.shared import Pt, Mm, RGBColor
from docx.enum.dml import MSO_COLOR_TYPE
def doctable(data, tabletitle, pathfile):
    document = Document()

    section = document.sections[0]
    section.page_height = Mm(297)
    section.page_width = Mm(210)
    section.left_margin = Mm(20)
    section.right_margin = Mm(20)
    section.top_margin = Mm(20)
    section.bottom_margin = Mm(20)
    section.header_distance = Mm(12.7)
    section.footer_distance = Mm(12.7)
    data = pd.DataFrame(data) # My input data is in the 2D list form
    head = document.add_heading(tabletitle)    
    head.style.font.name = 'Times New Roman'
    head.style.font.size =Pt(8)
    head.style.font.bold =True
    #head.style.font.color.rgb =RGBColor(0,0,0)
    table = document.add_table(rows=(data.shape[0]), cols=data.shape[1]) # First row are table headers!
    table.style.font.name = 'Times New Roman'
    table.style.font.size =Pt(8)
    table.style.font.bold =False
    table.allow_autofit = True
    table.autofit = True
    for i,c in enumerate(data.columns):
        table.cell(0,i).text = str(c)
        for row in range(1,data.shape[0]):
            table.cell(row, i).text = str(data[c][row]).replace('.','·')
    document.save(pathfile)
    return 0

# %%
from glob import glob
l = glob('../bivariate/pretty*')
for f in l:
    name = f.split('.')[-2].split('pretty_')[1]
    s = name.split('_')
    test,variable=tdict.get(s[0]),vdict.get(s[1])
    data = pd.read_csv(f,index_col=None)
    doctable(data,"Comparisons with {}, {}".format(variable,test),'../docxtables/{}.docx'.format(name))
# %%
d = pd.read_csv('../bivariate/pretty_fish_24days.csv',index_col=None)
doctable(d,"Death < 24 days, Fisher's exact tests",'../docxtables/fish_24days.docx')
# %%
vdict = {
    '24days': 'Death < 24 days',
    'month': 'Death < 28 days',
    'deathfree': 'Death-free days',
    'mmi': 'MMI',
    'mmi_bin': 'MMI > 0',
}
tdict= {
    'fish': "Fisher's exact test.",
    'ranksums': "Wilcoxon rank sums test.",
    'spear': "correlation test.",
}
# %%
